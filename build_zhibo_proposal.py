from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\yutiancheng\Documents\智博罗盘\智博罗盘申博一站式服务APP商业功能开发方案_升级版.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(90, 98, 110)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
WHITE = "FFFFFF"
BORDER = "D7DBE2"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        tbl_grid.append(col)
    set_table_borders(table)


def clear_para(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def add_para(doc, text="", style=None, size=None, color=INK, bold=False, italic=False, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    clear_para(cell.paragraphs[0])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows, widths=(1.55, 4.95), header=None):
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    set_table_width(table, list(widths))
    ridx = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        shade_cell(table.cell(0, 0), BLUE_FILL)
        p = table.cell(0, 0).paragraphs[0]
        clear_para(p)
        r = p.add_run(header)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        ridx = 1
    for label, value in rows:
        c0, c1 = table.row_cells(ridx)
        shade_cell(c0, LIGHT_FILL)
        for cell in (c0, c1):
            clear_para(cell.paragraphs[0])
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, size=10, color=DARK_BLUE, bold=True)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, size=10, color=INK)
        ridx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, widths)
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, BLUE_FILL)
        clear_para(cell.paragraphs[0])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for ridx, row in enumerate(rows, 1):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            clear_para(cell.paragraphs[0])
            p = cell.paragraphs[0]
            if cidx in (0, len(row) - 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=9.2, color=INK, bold=(cidx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


modules = [
    {
        "name": "1. 全域高校导师数据库",
        "tag": "核心基础板块",
        "goal": "建立可检索、可追溯、可持续更新的全国高校博士导师与院系招生信息底座，为AI匹配、学生投递、导师入驻、社区内容和商业化服务提供统一数据源。",
        "users": "学生用户、平台运营、入驻导师/院系管理员",
        "functions": [
            "三级目录：全国高校、二级学院、导师/课题组，支持按地区、学科门类、院校层次、招生方式逐层浏览。",
            "导师详情页：展示研究方向、近三年论文/项目、招生状态、招生偏好、官方联系方式、在读人数、毕业要求与可追溯来源链接。",
            "院校招生信息页：汇总招生简章、申请制/统考要求、时间节点、材料清单、学费学制、历年报录比与官方公告。",
            "多维搜索筛选：支持姓名、关键词、研究方向、地区、招收在职博士、是否接受跨专业、近期名额、申请截止时间等组合检索。",
            "数据更新后台：运营端录入、批量导入、来源校验、变更记录、过期提醒和版本回滚。",
        ],
        "data": [
            "导师实体：姓名、院校、学院、学科、研究方向、邮箱/主页、招生状态、名额、偏好标签、来源URL、更新时间。",
            "院校实体：招生年度、申请时间、材料要求、考试/面试要求、费用、公告链接、订阅人数。",
            "质量字段：来源可信度、人工复核状态、用户纠错次数、最近同步时间。",
        ],
        "workflow": [
            "运营导入或同步高校官网数据；系统生成待复核条目。",
            "人工审核关键字段后发布至前端；用户可收藏、订阅、纠错。",
            "数据变更触发目标用户推送，并同步影响AI匹配结果。",
        ],
        "business": "导师数据库是平台流量入口和会员转化基础。基础查询免费，高级筛选、批量导出、目标院校订阅、历史名额趋势可纳入AI会员或资料会员权益。",
        "priority": "一期必须上线基础数据、搜索、导师详情、运营后台；二期补齐订阅、纠错、数据可信度评分；三期做趋势分析与智能推荐。",
    },
    {
        "name": "2. AI智能导师匹配系统",
        "tag": "AI核心特色板块",
        "goal": "用AI解析学生背景并与导师数据库进行多维匹配，替代人工海量筛选，输出可解释、可执行、可持续更新的择校择导方案。",
        "users": "学生用户、付费会员、平台顾问",
        "functions": [
            "简历/RP/论文列表解析：支持PDF、Word、手动录入，抽取院校、专业、GPA、论文、项目、技能、研究兴趣等结构化字段。",
            "匹配分层：输出冲刺、稳申、保底三类导师清单，并给出匹配度、风险点、推荐理由和补强建议。",
            "研究方向语义匹配：不只匹配关键词，还识别研究主题、方法论、学科交叉和导师近期论文主题。",
            "陶瓷清单生成：自动生成导师联系人、优先级、沟通切入点、套磁风险提醒和跟进时间表。",
            "会员增值：解锁无限次解析、多版本简历对比、专属陶瓷邮件草稿、目标院校包和顾问复核。",
        ],
        "data": [
            "用户画像：教育背景、科研成果、目标地区、预算、申请方式、研究主题、时间规划。",
            "匹配特征：方向重合度、导师招生缺口、背景偏好、跨专业友好度、申请门槛、历史反馈。",
            "结果留痕：每次匹配版本、输入材料版本、推荐清单、用户操作记录。",
        ],
        "workflow": [
            "用户上传材料并确认AI抽取结果；系统要求用户补充目标偏好。",
            "AI调用导师数据库生成候选集，按分层策略输出名单。",
            "用户收藏/排除导师后，系统持续学习偏好并在数据更新时重新排序。",
        ],
        "business": "免费版提供有限匹配和基础解释；会员版提供深度匹配、无限刷新、导出清单、邮件生成和人工顾问复核。该模块是平台最强付费转化入口。",
        "priority": "一期上线简历解析和基础匹配；二期上线分层推荐、可解释评分和陶瓷清单；三期上线多轮对话式规划和顾问复核闭环。",
    },
    {
        "name": "3. 导师招生名额入驻&学生投递",
        "tag": "师生双向对接板块",
        "goal": "让真实导师、课题组或院系可以发布招生需求，让学生在平台内完成材料投递、状态跟踪和消息沟通，形成可闭环的招生撮合链路。",
        "users": "入驻导师、院系招生负责人、学生用户、平台审核员",
        "functions": [
            "导师/院系入驻：提交工牌、官网任职页、院系邮箱或单位证明，经平台审核后开通发布权限。",
            "招生名额发布：填写名额数量、研究方向、申请条件、材料清单、截止时间、课题内容、经费/实验室资源。",
            "学生一键投递：从个人材料库选择简历、RP、成绩单、论文、推荐信等材料并提交。",
            "投递状态流转：待查看、已查看、材料补充、面试邀请、暂不匹配、拟录取、已结束。",
            "内置消息沟通：支持模板回复、材料补交提醒、面试安排、沟通留痕与举报入口。",
        ],
        "data": [
            "招生职位：导师ID、招生年度、方向、名额、截止时间、要求、状态、浏览量、投递量。",
            "投递记录：学生ID、材料版本、投递时间、状态、导师反馈、消息记录。",
            "风控字段：资质审核记录、异常投递频率、投诉记录、联系方式展示策略。",
        ],
        "workflow": [
            "导师提交资质；平台审核通过后发布招生名额。",
            "学生浏览信息并投递材料；导师在后台筛选、回复、更新状态。",
            "系统将投递状态同步到学生申请进度中心，并沉淀匿名化转化数据。",
        ],
        "business": "基础发布可免费以吸引供给侧入驻；置顶展示、精准推荐、招生数据看板、院系品牌页可作为导师端或院系端付费能力。",
        "priority": "一期实现入驻审核、招生发布、一键投递；二期补齐消息、状态追踪和数据看板；三期增加机构级院系后台和推广产品。",
    },
    {
        "name": "4. 申博垂直论坛社区",
        "tag": "内容与信任沉淀板块",
        "goal": "建设只围绕博士申请的高信噪比社区，沉淀经验、问答、避坑、导师口碑和真实案例，降低新手信息成本并提升平台留存。",
        "users": "学生用户、上岸学长学姐、入驻导师、平台运营",
        "functions": [
            "精细分区：院校经验、导师陶瓷、RP写作、面试真题、避坑指南、跨专业申博、在职申博等。",
            "基础互动：发帖、评论、点赞、收藏、转发、匿名发帖、关注话题、站内提醒。",
            "经验沉淀：精华帖、合集、FAQ、年度申请复盘、院校专题页。",
            "导师风评：匿名评价导师指导风格、组会频率、毕业要求、科研氛围和招生真实性，需审核后展示。",
            "问答互助：免费问答和付费问答并存，认证博士、导师、顾问可获得答主身份。",
        ],
        "data": [
            "内容字段：话题、院校、学科、导师、申请年份、匿名标识、审核状态、热度。",
            "社区画像：用户关注方向、收藏内容、互动行为、贡献等级、信用分。",
            "风评字段：评价维度、证据材料、时间、审核结论、争议处理记录。",
        ],
        "workflow": [
            "用户发帖或匿名评价；AI初审识别敏感、造谣、广告和隐私泄露风险。",
            "运营复核高风险内容；通过后进入对应分区和导师/院校关联页。",
            "优质内容被收录为专题或FAQ，并反哺导师数据库和申请工具箱。",
        ],
        "business": "社区本身不应强付费，重点承担获客、留存、信任和内容资产沉淀。可通过付费问答、资料包导流、会员权益和RP辅导转化实现间接商业化。",
        "priority": "一期上线分区、发帖、评论、收藏；二期上线匿名、精华、问答；三期上线导师风评体系和内容推荐。",
    },
    {
        "name": "5. RP书写付费辅导商业化",
        "tag": "核心营收板块",
        "goal": "围绕博士申请中最难标准化、最具付费意愿的Research Proposal，建立AI预审、套餐购买、老师接单、交付验收、评价售后的商业闭环。",
        "users": "学生用户、辅导老师、平台客服/售后、平台财务",
        "functions": [
            "AI免费预审：检测主题清晰度、研究问题、创新性、文献结构、方法可行性、语言表达和院校适配度。",
            "标准化套餐：基础润色、框架指导、全文修改、选题+写作+修改、院校专属RP定制。",
            "老师入驻审核：验证博士身份、高校经历、论文成果、辅导案例、擅长学科和服务承诺。",
            "订单履约：下单、材料上传、老师接单、阶段交付、修改次数、交付确认、退款/仲裁。",
            "付费解锁联系方式：购买后可按规则解锁微信或专属联系方式，同时保留平台内订单留痕。",
        ],
        "data": [
            "商品字段：套餐类型、价格、服务范围、交付周期、修改次数、适用学科、退款规则。",
            "老师字段：资质、擅长领域、评分、接单量、响应速度、历史纠纷率。",
            "订单字段：材料版本、阶段状态、交付文件、沟通记录、评价、佣金结算。",
        ],
        "workflow": [
            "用户先做AI预审，系统提示可购买的服务层级。",
            "用户选择套餐和老师，下单后上传材料并进入订单履约。",
            "老师交付阶段成果；用户确认或申请修改；平台完成分账和评价沉淀。",
        ],
        "business": "平台核心收入来自订单抽佣、老师服务费、加急费、院校定制包和高阶全流程托管。要用标准化边界控制交付风险，避免无边界承诺。",
        "priority": "一期上线AI预审、套餐、下单、老师审核和订单管理；二期上线售后仲裁、评价和分账；三期上线全流程托管和企业化顾问后台。",
    },
    {
        "name": "6. 申博工具箱&备考资讯",
        "tag": "高频使用与留存板块",
        "goal": "把申博过程中分散在官网、公众号、网盘和社群里的资讯、模板和备考工具集中到一个可订阅、可提醒、可复用的工具中心。",
        "users": "学生用户、平台运营、内容合作方",
        "functions": [
            "资讯中枢：招生简章、政策变动、报名时间、直博/硕博连读通知、目标院校订阅和推送。",
            "陶瓷工具：邮件模板、AI润色、跟进提醒、导师背景摘要、常见失误检查。",
            "材料模板库：简历、推荐信、个人陈述、RP结构模板、英文邮件模板。",
            "备考资源库：考博英语、专业课真题、面试题、知识点笔记、模拟题库。",
            "面试训练：自我介绍模板、高频问题、院校面试风格、模拟面试记录。",
        ],
        "data": [
            "资讯字段：院校、学院、学科、年份、截止时间、来源、推送状态。",
            "模板字段：适用场景、语言、学科、下载权限、使用次数、评分。",
            "资源字段：真题年份、科目、来源、版权/授权状态、付费权限。",
        ],
        "workflow": [
            "运营发布或更新资讯；订阅用户收到提醒。",
            "用户从工具箱生成邮件、材料或面试计划，并保存到个人中心。",
            "使用行为反哺AI匹配和申请进度提醒。",
        ],
        "business": "基础资讯免费拉新；高价值资料包、真题库、模板下载、AI邮件润色次数、模拟面试可作为会员或单次付费产品。",
        "priority": "一期上线资讯、模板和邮件工具；二期上线订阅推送、真题库和面试题；三期上线智能备考计划。",
    },
    {
        "name": "7. 个人申请进度管理中心",
        "tag": "长期留存与复购板块",
        "goal": "为多院校、多导师、多材料版本的博士申请过程建立个人台账，帮助用户按时间、材料、投递和结果管理完整申请周期。",
        "users": "学生用户、顾问服务人员",
        "functions": [
            "申请台账：记录目标院校、导师、投递日期、材料状态、审核进度、面试时间、录取结果。",
            "材料云档案：管理简历、RP、成绩单、推荐信、证书等版本，支持一键复用到投递。",
            "智能提醒：截止时间、补材料、导师跟进、面试倒计时、结果确认。",
            "规划日历：按月份展示申博关键任务，并支持导出个人计划。",
            "复盘看板：统计投递数量、回复率、面试率、录取结果和下一步建议。",
        ],
        "data": [
            "申请记录：目标、状态、优先级、时间节点、材料版本、备注。",
            "提醒字段：提醒类型、触发时间、关联导师/院校、完成状态。",
            "材料字段：文件类型、版本、适用院校、最近修改时间、隐私权限。",
        ],
        "workflow": [
            "用户从导师库、AI匹配或招生投递页加入目标到台账。",
            "系统自动生成节点和提醒；用户更新状态或由投递系统同步状态。",
            "申请结束后形成复盘数据，并推荐后续辅导、面试或补申服务。",
        ],
        "business": "基础台账免费；多材料版本管理、高级提醒、批量导出、顾问协作视图和申请复盘报告可纳入会员权益。",
        "priority": "一期上线台账、材料库、基础提醒；二期同步投递状态和日历；三期上线复盘报告和顾问协作。",
    },
    {
        "name": "8. 平台审核&诚信体系",
        "tag": "风控合规底座",
        "goal": "通过身份认证、资质审核、内容审核、交易保障和信用评价，降低虚假导师、虚假招生、诈骗、学术不端和服务纠纷风险。",
        "users": "全体用户、平台审核员、客服/法务",
        "functions": [
            "双向实名认证：学生手机号认证；导师、院系、辅导老师需提交可核验证明。",
            "内容审核：AI+人工审核招生信息、社区帖子、导师风评、付费服务描述和私信高风险内容。",
            "信用评分：结合真实资料、历史行为、投诉、履约、评价和违规记录生成信用标签。",
            "举报维权：一键举报、证据上传、处理进度、处罚结果通知。",
            "交易保障：订单留痕、交付验收、退款规则、争议仲裁、黑名单机制。",
        ],
        "data": [
            "认证记录：认证方式、证明材料、审核人、审核时间、过期时间。",
            "风控记录：违规类型、证据、处理状态、处罚措施、申诉结果。",
            "信用记录：评分维度、加减分事件、展示标签、冻结/解封状态。",
        ],
        "workflow": [
            "用户提交认证或内容；系统预审并分配风险等级。",
            "低风险自动通过，高风险进入人工复核和补充材料流程。",
            "违规处理结果影响账号权限、内容展示、交易资格和推荐权重。",
        ],
        "business": "诚信体系本身不是直接收入模块，但决定平台能否长期商业化。它保护付费订单、导师入驻、社区风评和品牌可信度，是商业闭环的底层保险。",
        "priority": "一期上线身份审核、内容审核和举报；二期上线信用分、交易仲裁；三期上线风控模型和合规审计报表。",
    },
]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    clear_para(header)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("智博罗盘 | 申博一站式服务APP商业功能开发方案")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    clear_para(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Confidential Draft")
    set_run_font(r, size=9, color=MUTED)

    add_para(doc, "智博罗盘产品方案", size=10.5, color=MUTED, bold=True, after=4)
    add_para(doc, "申博一站式服务APP", size=24, color=INK, bold=True, after=2)
    add_para(doc, "商业功能开发方案（升级版）", size=17, color=DARK_BLUE, bold=True, after=12)
    add_kv_table(
        doc,
        [
            ("文档用途", "用于产品立项、技术开发拆解、商业化评审与后续PRD细化。"),
            ("核心定位", "AI赋能、导师数据库驱动、师生投递闭环、社区信任沉淀、RP辅导商业化的垂直申博平台。"),
            ("版本日期", f"{date.today().isoformat()}"),
            ("建议读者", "创始团队、产品经理、技术负责人、运营负责人、商业合作方。"),
        ],
        header="文档摘要",
    )
    add_callout(
        doc,
        "一句话定位",
        "智博罗盘不是单一资讯站或社区，而是围绕博士申请全过程建立“信息查询-AI匹配-招生投递-社区验证-材料辅导-进度管理-风控审核”的闭环产品。",
    )

    add_h1(doc, "一、产品定位与建设原则")
    add_para(
        doc,
        "本方案将原有8大功能板块升级为可开发、可运营、可商业化的产品架构。产品面向国内硕博在读学生、应届硕士、在职申博人群，同时服务高校博导、院系招生负责人、RP辅导老师和平台运营团队。",
    )
    add_h2(doc, "1.1 核心痛点")
    for item in [
        "申博信息分散在高校官网、论坛、公众号、社群和个人经验贴中，用户需要跨平台搜集，效率低且可信度不稳定。",
        "导师研究方向、招生状态、招生偏好和真实口碑缺乏结构化沉淀，导致用户陶瓷和投递方向模糊。",
        "RP写作、套磁邮件、面试准备等关键环节需要专业指导，但市场服务分散、价格不透明、售后难保障。",
        "多院校、多导师、多材料版本并行推进时，用户容易遗漏截止时间、材料状态和沟通跟进。",
    ]:
        add_bullet(doc, item)
    add_h2(doc, "1.2 产品原则")
    for item in [
        "先建立可信数据底座，再做AI推荐和商业化转化。",
        "所有AI结果必须可解释、可追溯、可让用户修正，不做黑箱式推荐。",
        "社区承担信任沉淀，不以牺牲内容真实性换取短期流量。",
        "付费服务要标准化边界、过程留痕、交付可验收，避免变成不可控的私域交易。",
        "审核与诚信体系从一期开始建设，避免平台在规模化后再补风控。",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "二、8大核心功能架构总览")
    add_matrix(
        doc,
        ["板块", "核心价值", "一期交付重点", "商业化属性"],
        [
            ("导师数据库", "平台数据底座与流量入口", "高校/学院/导师目录、搜索、详情页", "会员筛选、订阅、数据服务"),
            ("AI匹配", "核心差异化与付费转化", "简历解析、基础推荐、匹配说明", "AI会员、顾问复核"),
            ("招生入驻&投递", "师生双向撮合闭环", "导师审核、招生发布、一键投递", "导师曝光、院系服务"),
            ("垂直社区", "内容沉淀与信任验证", "分区、发帖、评论、收藏", "问答、资料导流、服务转化"),
            ("RP辅导", "核心现金流", "AI预审、套餐、订单、老师审核", "抽佣、加急、托管"),
            ("工具箱&资讯", "高频留存与材料复用", "资讯、模板、邮件工具", "资料包、模板会员"),
            ("申请进度中心", "长期留存与复购触发", "台账、材料库、提醒", "高级提醒、顾问协作"),
            ("审核诚信体系", "平台安全与品牌信任", "认证、内容审核、举报", "间接保障交易与供给质量"),
        ],
        [1.28, 1.95, 1.95, 1.32],
    )

    add_h1(doc, "三、各板块详细功能开发细则")
    for module in modules:
        add_h2(doc, f"{module['name']}（{module['tag']}）")
        add_kv_table(
            doc,
            [
                ("模块目标", module["goal"]),
                ("主要用户", module["users"]),
                ("商业定位", module["business"]),
                ("迭代优先级", module["priority"]),
            ],
            widths=(1.35, 5.15),
        )
        add_h3(doc, "核心功能")
        for item in module["functions"]:
            add_bullet(doc, item)
        add_h3(doc, "关键数据字段")
        for item in module["data"]:
            add_bullet(doc, item)
        add_h3(doc, "推荐业务流程")
        for item in module["workflow"]:
            add_number(doc, item)

    add_h1(doc, "四、用户角色与权限划分")
    add_matrix(
        doc,
        ["角色", "核心权限", "关键限制", "运营价值"],
        [
            ("学生用户", "查询导师、AI匹配、社区互动、投递材料、购买辅导、管理进度", "敏感评价需审核；高频投递需风控", "主要流量与付费来源"),
            ("入驻导师/院系", "发布招生、查看材料、沟通学生、更新状态、查看数据看板", "需资质审核；禁止虚假宣传和私下违规收费", "增加真实供给与平台权威性"),
            ("辅导老师", "上架RP服务、接单、交付、售后沟通、查看评价", "需学历/成果审核；订单需平台留痕", "核心商业化供给"),
            ("平台管理员", "数据维护、资质审核、内容审核、订单仲裁、风控处理、运营推荐", "需分级权限和操作日志", "保障合规、质量和增长"),
        ],
        [1.25, 2.55, 1.55, 1.15],
    )

    add_h1(doc, "五、商业化体系设计")
    add_para(doc, "商业化应围绕用户真实刚需分层设计，避免过早把基础信息全部收费。建议采用“免费入口+AI会员+付费辅导+资料服务+导师端增值”的组合模型。")
    for item in [
        "核心收入：RP辅导订单抽佣、全流程申博托管、加急修改、院校定制RP服务。",
        "会员收入：AI高级匹配、无限简历解析、匹配清单导出、目标院校订阅、高级提醒和申请复盘报告。",
        "资料收入：真题库、稀缺经验包、模板包、面试题库和院校专题资料。",
        "导师/院系端收入：招生信息置顶、精准曝光、投递数据看板、院系品牌页和批量筛选工具。",
        "顾问服务收入：择校规划、陶瓷策略、面试模拟、材料组合优化和年度申请陪跑。",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "六、技术开发迭代优先级")
    add_matrix(
        doc,
        ["阶段", "目标", "核心范围", "验收标准"],
        [
            ("一期", "跑通供需和现金流最小闭环", "导师数据库、搜索、基础AI匹配、招生发布/投递、社区基础、RP套餐订单、基础审核", "用户可查导师、获得推荐、投递招生、购买RP服务；后台可审核和维护数据"),
            ("二期", "提升匹配精度和留存", "AI分层匹配、陶瓷工具箱、资讯订阅、进度中心、私信、评价、售后仲裁", "用户可形成完整申请台账；导师与学生可在站内沟通；订单可闭环售后"),
            ("三期", "规模化商业化与数据智能", "会员体系、资料商城、院系端看板、风控模型、推荐算法、全流程托管", "平台具备稳定付费转化、供给质量控制和运营数据看板"),
        ],
        [0.85, 1.55, 2.55, 1.55],
    )

    add_h1(doc, "七、关键页面与后台清单")
    add_h2(doc, "7.1 用户端关键页面")
    for item in [
        "首页：搜索入口、AI匹配入口、最新招生、热门社区、RP预审入口。",
        "导师库：筛选页、导师详情页、院校详情页、收藏订阅页。",
        "AI匹配：材料上传、画像确认、匹配结果、陶瓷清单、导出与会员升级。",
        "招生投递：招生详情、材料选择、投递确认、状态跟踪、消息沟通。",
        "社区：分区列表、帖子详情、发帖页、匿名评价、问答页。",
        "RP辅导：AI预审、套餐列表、老师主页、订单详情、售后评价。",
        "个人中心：申请台账、材料库、日历提醒、会员权益、订单与钱包。",
    ]:
        add_bullet(doc, item)
    add_h2(doc, "7.2 后台关键页面")
    for item in [
        "数据后台：高校、学院、导师、招生简章、来源链接、更新记录。",
        "审核后台：导师/老师资质、社区内容、招生信息、举报与申诉。",
        "订单后台：RP订单、交付节点、退款仲裁、分账结算。",
        "运营后台：Banner、专题、精华帖、资料包、推送、用户分群。",
        "风控后台：异常账号、投诉记录、信用分、黑名单、操作日志。",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "八、合规与风险控制")
    for item in [
        "导师风评必须设置审核、申诉和证据机制，避免侵犯名誉权或传播未经核实的信息。",
        "付费辅导必须明确服务边界，禁止代写论文、伪造材料、保录承诺等高风险宣传。",
        "用户上传材料涉及个人隐私和学术成果，应提供权限控制、加密存储、删除机制和隐私政策说明。",
        "高校官网数据需保留来源链接和更新时间，避免平台展示过期招生信息造成用户损失。",
        "AI推荐需要提示“仅供参考”，并展示影响推荐的关键因素，降低误导性决策风险。",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "九、阶段性验收指标")
    add_matrix(
        doc,
        ["维度", "一期建议指标", "二期建议指标", "三期建议指标"],
        [
            ("数据覆盖", "覆盖重点高校与核心学院，导师信息可检索", "覆盖主要学科门类，信息有更新机制", "建立数据可信度评分和趋势分析"),
            ("AI效果", "基础匹配可用，用户可理解推荐理由", "分层匹配稳定，收藏/排除可反馈", "推荐与商业服务转化联动"),
            ("交易闭环", "RP订单可下单、接单、交付、评价", "售后仲裁和分账稳定", "托管服务和会员复购成熟"),
            ("社区质量", "分区清晰，内容可审核", "精华沉淀和问答机制可用", "导师口碑与院校专题形成资产"),
            ("风控合规", "身份审核、举报、内容审核上线", "信用体系和交易仲裁上线", "风控模型和审计报表上线"),
        ],
        [1.1, 1.8, 1.8, 1.8],
    )

    add_h1(doc, "十、结论")
    add_para(
        doc,
        "智博罗盘的产品价值不在于单点功能，而在于把申博过程中的信息、匹配、投递、社区、材料辅导和进度管理连接成一条可持续运营的闭环。建议一期优先完成导师数据库、基础AI匹配、招生投递、RP订单和审核后台，让平台尽快具备真实使用场景与商业化验证能力；二期再强化留存与信任；三期推进会员、数据智能和规模化商业化。",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
